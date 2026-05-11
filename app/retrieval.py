import os
import re
import sys
import csv
import uuid
import json
import shutil
import hashlib
import asyncio

# Fix Windows cp1252 emoji encoding
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from pypdf import PdfReader

from sklearn.metrics.pairwise import cosine_similarity

from langchain_community.docstore.document import Document

from langchain_community.vectorstores import FAISS

from langchain_huggingface import HuggingFaceEmbeddings

from langchain_text_splitters import RecursiveCharacterTextSplitter


# =====================================================
# OPTIONAL IMPORTS
# =====================================================

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


# =====================================================
# RETRIEVAL ENGINE
# =====================================================

class RetrievalEngine:

    def __init__(self):

        # =============================================
        # EMBEDDINGS
        # =============================================

        print("\n🔄 Loading Embeddings...\n")

        self.embedding_model = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            encode_kwargs={"normalize_embeddings": True}
        )

        print("\n✅ Embeddings Loaded\n")

        # =============================================
        # TEXT SPLITTER
        # =============================================

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=700,
            chunk_overlap=120,
            separators=[
                "\n# ", "\n## ", "\n### ",
                "\n\n", "\n",
                ". ", "? ", "! ", "; ",
                " ", ""
            ]
        )

        # =============================================
        # PATHS
        # =============================================

        self.faiss_path = "data/faiss_index"

        os.makedirs("data", exist_ok=True)

        # =============================================
        # VECTOR STORE + ACTIVE DOCS
        # =============================================

        self.pdf_vectordb = None
        self.active_documents = []

    # =================================================
    # CLEAN TEXT
    # =================================================

    def clean_text(self, text):
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"\n+", "\n", text)
        return text.strip()

    # =================================================
    # HASH
    # =================================================

    def create_hash(self, text):
        return hashlib.md5(text.encode()).hexdigest()

    # =================================================
    # DEDUPLICATION
    # =================================================

    def deduplicate_chunks(self, chunks):
        unique = []
        seen = set()
        for chunk in chunks:
            chunk_hash = self.create_hash(chunk["content"])
            if chunk_hash in seen:
                continue
            seen.add(chunk_hash)
            unique.append(chunk)
        return unique

    # =================================================
    # UNIVERSAL FILE PROCESSOR
    # =================================================

    def process_file(self, file_path):
        """
        Dispatch to the correct handler based on file extension.
        Supports: PDF, TXT, MD, CSV, JSON, DOCX
        """
        ext = os.path.splitext(file_path)[1].lower()

        print(f"\n📂 Processing file [{ext}]: {file_path}\n")

        if ext == ".pdf":
            return self._process_pdf(file_path)
        elif ext in (".txt", ".md"):
            return self._process_text(file_path)
        elif ext == ".csv":
            return self._process_csv(file_path)
        elif ext == ".json":
            return self._process_json(file_path)
        elif ext == ".docx":
            return self._process_docx(file_path)
        else:
            # Fallback: try reading as plain text
            return self._process_text(file_path)

    # =================================================
    # PDF PROCESSOR
    # =================================================

    def _process_pdf(self, file_path):
        print(f"\n📄 Processing PDF: {file_path}\n")

        reader = PdfReader(file_path)
        print(f"\n📄 Total Pages: {len(reader.pages)}\n")

        document_name = os.path.basename(file_path)
        document_id = str(uuid.uuid4())
        docs = []
        total_chunks = 0

        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
            except Exception as e:
                print(f"\n❌ Page Error: {e}")
                continue

            if not page_text:
                continue

            page_text = self.clean_text(page_text)

            if len(page_text) < 50:
                continue

            chunks = self.text_splitter.split_text(page_text)
            print(f"📄 Page {page_num+1} Chunks: {len(chunks)}")
            total_chunks += len(chunks)

            for chunk_idx, chunk in enumerate(chunks):
                docs.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "document_name": document_name,
                        "document_id": document_id,
                        "page_number": page_num + 1,
                        "chunk_index": chunk_idx,
                        "chunk_id": f"{document_id}_{page_num}_{chunk_idx}",
                        "type": "pdf",
                        "retrieval_weight": 1.5
                    }
                ))

        return self._index_documents(docs, document_name, total_chunks)

    # =================================================
    # TEXT / MARKDOWN PROCESSOR
    # =================================================

    def _process_text(self, file_path):
        document_name = os.path.basename(file_path)
        document_id = str(uuid.uuid4())

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
        except Exception as e:
            print(f"\n❌ Text Read Error: {e}")
            return

        if not raw_text.strip():
            print("\n❌ Empty file\n")
            return

        raw_text = self.clean_text(raw_text)
        chunks = self.text_splitter.split_text(raw_text)
        docs = []

        for chunk_idx, chunk in enumerate(chunks):
            docs.append(Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "document_name": document_name,
                    "document_id": document_id,
                    "page_number": None,
                    "chunk_index": chunk_idx,
                    "chunk_id": f"{document_id}_{chunk_idx}",
                    "type": "text",
                    "retrieval_weight": 1.5
                }
            ))

        return self._index_documents(docs, document_name, len(chunks))

    # =================================================
    # CSV PROCESSOR
    # =================================================

    def _process_csv(self, file_path):
        document_name = os.path.basename(file_path)
        document_id = str(uuid.uuid4())
        docs = []
        chunk_idx = 0

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.DictReader(f)
                rows = list(reader)

            # Convert rows to text blocks of 20 rows each
            batch_size = 20
            for i in range(0, len(rows), batch_size):
                batch = rows[i:i + batch_size]
                text_lines = []
                for row in batch:
                    line = " | ".join(
                        f"{k}: {v}" for k, v in row.items() if v
                    )
                    text_lines.append(line)
                block = "\n".join(text_lines)
                block = self.clean_text(block)

                if len(block) < 30:
                    continue

                docs.append(Document(
                    page_content=block,
                    metadata={
                        "source": file_path,
                        "document_name": document_name,
                        "document_id": document_id,
                        "page_number": None,
                        "chunk_index": chunk_idx,
                        "chunk_id": f"{document_id}_{chunk_idx}",
                        "type": "csv",
                        "retrieval_weight": 1.5
                    }
                ))
                chunk_idx += 1

        except Exception as e:
            print(f"\n❌ CSV Error: {e}")
            return

        return self._index_documents(docs, document_name, len(docs))

    # =================================================
    # JSON PROCESSOR
    # =================================================

    def _process_json(self, file_path):
        document_name = os.path.basename(file_path)
        document_id = str(uuid.uuid4())

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                data = json.load(f)
            raw_text = json.dumps(data, indent=2)
        except Exception as e:
            print(f"\n❌ JSON Error: {e}")
            return

        raw_text = self.clean_text(raw_text)
        chunks = self.text_splitter.split_text(raw_text)
        docs = []

        for chunk_idx, chunk in enumerate(chunks):
            docs.append(Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "document_name": document_name,
                    "document_id": document_id,
                    "page_number": None,
                    "chunk_index": chunk_idx,
                    "chunk_id": f"{document_id}_{chunk_idx}",
                    "type": "json",
                    "retrieval_weight": 1.5
                }
            ))

        return self._index_documents(docs, document_name, len(chunks))

    # =================================================
    # DOCX PROCESSOR
    # =================================================

    def _process_docx(self, file_path):
        if not DOCX_AVAILABLE:
            print("\n❌ python-docx not installed. Run: pip install python-docx\n")
            return

        document_name = os.path.basename(file_path)
        document_id = str(uuid.uuid4())

        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            raw_text = "\n".join(paragraphs)
        except Exception as e:
            print(f"\n❌ DOCX Error: {e}")
            return

        raw_text = self.clean_text(raw_text)
        chunks = self.text_splitter.split_text(raw_text)
        docs = []

        for chunk_idx, chunk in enumerate(chunks):
            docs.append(Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "document_name": document_name,
                    "document_id": document_id,
                    "page_number": None,
                    "chunk_index": chunk_idx,
                    "chunk_id": f"{document_id}_{chunk_idx}",
                    "type": "docx",
                    "retrieval_weight": 1.5
                }
            ))

        return self._index_documents(docs, document_name, len(chunks))

    # =================================================
    # INDEX DOCUMENTS (shared indexing logic)
    # =================================================

    def _index_documents(self, docs, document_name, total_chunks):
        if not docs:
            print("\n❌ No Chunks Created\n")
            return

        # Remove old FAISS index
        if os.path.exists(self.faiss_path):
            print("\n🗑️ Removing Old FAISS Index...\n")
            shutil.rmtree(self.faiss_path)

        # Reset memory
        self.pdf_vectordb = None
        self.active_documents = []

        print("\n🔄 Creating Fresh FAISS Index...\n")

        self.pdf_vectordb = FAISS.from_documents(docs, self.embedding_model)

        print(f"\n✅ VECTOR COUNT: {self.pdf_vectordb.index.ntotal}\n")

        self.pdf_vectordb.save_local(self.faiss_path)

        self.active_documents = [document_name]

        print(f"\n✅ Indexed: {document_name}")
        print(f"📄 Total Chunks: {total_chunks}\n")

    # =================================================
    # KEYWORD OVERLAP SCORE
    # =================================================

    def keyword_overlap_score(self, query, content):
        query_words = set(query.lower().split())
        content_words = set(content.lower().split())
        overlap = len(query_words.intersection(content_words))
        return overlap * 0.01

    # =================================================
    # FORMAT RESULTS
    # =================================================

    def format_results(self, query, results):
        formatted = []

        for doc, raw_score in results:

            groundedness = max(0.0, round(1 - float(raw_score), 3))

            retrieval_weight = doc.metadata.get("retrieval_weight", 1.0)

            overlap_score = self.keyword_overlap_score(
                query, doc.page_content
            )

            final_score = round(
                (groundedness * retrieval_weight) + overlap_score,
                3
            )

            # Relaxed threshold for better recall
            if final_score < 0.01:
                continue

            formatted.append({
                "content": doc.page_content,
                "source": doc.metadata.get("source"),
                "document_name": doc.metadata.get("document_name"),
                "document_id": doc.metadata.get("document_id"),
                "page_number": doc.metadata.get("page_number"),
                "chunk_index": doc.metadata.get("chunk_index"),
                "chunk_id": doc.metadata.get("chunk_id"),
                "type": doc.metadata.get("type"),
                "score": final_score
            })

        formatted = self.deduplicate_chunks(formatted)

        formatted = sorted(
            formatted,
            key=lambda x: x["score"],
            reverse=True
        )

        return formatted

    # =================================================
    # DOCUMENT RETRIEVAL
    # =================================================

    def retrieve_from_docs(self, query, top_k=8):
        """
        Retrieve from indexed documents (PDF, TXT, DOCX, CSV, JSON).
        """
        try:
            if self.pdf_vectordb is None:
                print("\n❌ FAISS Not Loaded\n")
                return []

            print(f"\n🔍 Retrieval Query: {query}\n")
            print(f"\n📚 Active Docs: {self.active_documents}\n")

            results = self.pdf_vectordb.similarity_search_with_score(
                query, k=top_k
            )

            print(f"\n📄 RAW RESULTS: {len(results)}\n")

            for doc, score in results:
                print(f"\n📊 SCORE: {score}")
                print(f"📄 DOC: {doc.metadata.get('document_name')}")

            formatted = self.format_results(query, results)

            print(f"\n✅ FINAL RETRIEVED: {len(formatted)} chunks\n")

            return formatted

        except Exception as e:
            print(f"\n❌ DOC RETRIEVAL ERROR\n{e}")
            return []

    # Keep backward-compatible alias
    def retrieve_pdf(self, query, top_k=8):
        return self.retrieve_from_docs(query, top_k)

    # =================================================
    # WEB RETRIEVAL
    # =================================================

    def retrieve_web_documents(self, query, documents, top_k=5):
        docs = []

        for doc_idx, item in enumerate(documents):
            content = item.get("content", "")

            if len(content) < 300:
                continue

            content = self.clean_text(content)
            chunks = self.text_splitter.split_text(content)

            for chunk_idx, chunk in enumerate(chunks):
                docs.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": item.get("source", "web"),
                        "document_name": item.get("title", "web_document"),
                        "type": "web",
                        "chunk_index": chunk_idx,
                        "chunk_id": f"web_{doc_idx}_{chunk_idx}",
                        "retrieval_weight": 1.0
                    }
                ))

        if not docs:
            return []

        query_embedding = self.embedding_model.embed_query(query)
        doc_texts = [d.page_content for d in docs]
        doc_embeddings = self.embedding_model.embed_documents(doc_texts)

        similarities = cosine_similarity(
            [query_embedding], doc_embeddings
        )[0]

        ranked = sorted(
            zip(docs, similarities),
            key=lambda x: x[1],
            reverse=True
        )[:top_k]

        results = [(doc, 1 - score) for doc, score in ranked]

        return self.format_results(query, results)

    # =================================================
    # HYBRID FUSION
    # =================================================

    def fuse_contexts(self, doc_results, web_results):
        for item in doc_results:
            item["score"] *= 1.35

        combined = doc_results + web_results
        combined = self.deduplicate_chunks(combined)
        combined = sorted(combined, key=lambda x: x["score"], reverse=True)

        return combined[:15]

    # =================================================
    # BUILD CONTEXT
    # =================================================

    def build_context(self, retrieved_docs, max_chars=30000):
        context_parts = []
        current_length = 0

        for item in retrieved_docs:
            formatted = f"""
[SOURCE TYPE]: {item.get('type')}
[DOCUMENT]: {item.get('document_name')}
[PAGE]: {item.get('page_number', 'N/A')}
[GROUNDEDNESS]: {item.get('score')}

{item.get('content')}

"""
            if current_length + len(formatted) > max_chars:
                break

            context_parts.append(formatted)
            current_length += len(formatted)

        return "\n".join(context_parts)

    # =================================================
    # CLEAR DOCUMENTS  ← FIXED: properly inside class
    # =================================================

    def clear_documents(self):
        try:
            print("\n🗑️ Clearing Documents...\n")

            self.pdf_vectordb = None
            self.active_documents = []

            if os.path.exists(self.faiss_path):
                shutil.rmtree(self.faiss_path)

            print("\n✅ Documents Cleared\n")

        except Exception as e:
            print(f"\n❌ CLEAR DOCUMENT ERROR\n{e}")